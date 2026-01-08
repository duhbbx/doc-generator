"""Word document template renderer module."""

import re
from copy import deepcopy
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from .expression import ExpressionEvaluator


class WordRenderer:
    """Renders Word documents from templates with placeholder replacement."""

    PLACEHOLDER_PATTERN = re.compile(r"\{\{([^}]+)\}\}")

    def __init__(self, template_path: str | Path):
        """Initialize the Word renderer.

        Args:
            template_path: Path to the Word template file.
        """
        self.template_path = Path(template_path)
        self._document = None
        self._evaluator = ExpressionEvaluator()

    def load_template(self) -> None:
        """Load the Word template."""
        self._document = Document(self.template_path)

    def get_placeholders(self) -> list[str]:
        """Extract all placeholders from the template.

        Returns:
            List of unique placeholder names found in the template.
        """
        if not self._document:
            self.load_template()

        placeholders = set()

        # Search in paragraphs
        for para in self._document.paragraphs:
            text = self._get_paragraph_text(para)
            placeholders.update(self.PLACEHOLDER_PATTERN.findall(text))

        # Search in tables
        for table in self._document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = self._get_paragraph_text(para)
                        placeholders.update(self.PLACEHOLDER_PATTERN.findall(text))

        # Search in headers and footers
        for section in self._document.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        text = self._get_paragraph_text(para)
                        placeholders.update(self.PLACEHOLDER_PATTERN.findall(text))

            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        text = self._get_paragraph_text(para)
                        placeholders.update(self.PLACEHOLDER_PATTERN.findall(text))

        return sorted(placeholders)

    @staticmethod
    def _get_paragraph_text(paragraph: Paragraph) -> str:
        """Get full text from a paragraph, handling split runs."""
        return "".join(run.text for run in paragraph.runs)

    def _replace_in_paragraph(self, paragraph: Paragraph, data: dict[str, Any], mappings: dict[str, str]) -> None:
        """Replace placeholders in a paragraph while preserving formatting.

        Args:
            paragraph: The paragraph to process.
            data: Row data from Excel.
            mappings: Mapping from placeholder names to expressions.
        """
        # Get full paragraph text
        full_text = self._get_paragraph_text(paragraph)

        if not self.PLACEHOLDER_PATTERN.search(full_text):
            return

        # Find all placeholders in the paragraph
        for match in self.PLACEHOLDER_PATTERN.finditer(full_text):
            placeholder_name = match.group(1)
            placeholder_full = match.group(0)  # {{name}}

            # Get the expression for this placeholder
            expression = mappings.get(placeholder_name, f"{{{{{placeholder_name}}}}}")

            # Evaluate the expression
            try:
                value = self._evaluator.evaluate_safe(expression, data, default="")
                if value is None:
                    value = ""
                replacement = str(value)
            except Exception:
                replacement = ""

            # Replace in the text
            full_text = full_text.replace(placeholder_full, replacement, 1)

        # Now we need to update the paragraph runs with the new text
        # This is tricky because we want to preserve formatting
        self._update_paragraph_text(paragraph, full_text)

    def _update_paragraph_text(self, paragraph: Paragraph, new_text: str) -> None:
        """Update paragraph text while trying to preserve run formatting."""
        if not paragraph.runs:
            paragraph.add_run(new_text)
            return

        # Simple approach: put all text in the first run, clear the rest
        first_run = paragraph.runs[0]
        first_run.text = new_text

        for run in paragraph.runs[1:]:
            run.text = ""

    def _replace_in_table(self, table: Table, data: dict[str, Any], mappings: dict[str, str]) -> None:
        """Replace placeholders in a table."""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self._replace_in_paragraph(paragraph, data, mappings)

    def render(self, data: dict[str, Any], mappings: dict[str, str], output_path: str | Path) -> None:
        """Render the template with data and save to output path.

        Args:
            data: Row data from Excel (column name -> value).
            mappings: Mapping from placeholder names to expressions.
                      If a placeholder is not in mappings, it looks for
                      a direct match in data.
            output_path: Path to save the rendered document.
        """
        # Reload template for fresh copy
        document = Document(self.template_path)

        # Build complete mappings (add direct mappings for any unmapped placeholders)
        complete_mappings = dict(mappings)
        for key in data.keys():
            if key not in complete_mappings:
                complete_mappings[key] = f"{{{{{key}}}}}"

        # Replace in paragraphs
        for paragraph in document.paragraphs:
            self._replace_in_paragraph(paragraph, data, complete_mappings)

        # Replace in tables
        for table in document.tables:
            self._replace_in_table(table, data, complete_mappings)

        # Replace in headers and footers
        for section in document.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for paragraph in header.paragraphs:
                        self._replace_in_paragraph(paragraph, data, complete_mappings)

            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for paragraph in footer.paragraphs:
                        self._replace_in_paragraph(paragraph, data, complete_mappings)

        # Save the document
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)

    def generate_filename(self, pattern: str, data: dict[str, Any], index: int = 0) -> str:
        """Generate output filename from pattern and data.

        Args:
            pattern: Filename pattern with placeholders (e.g., "{{name}}_report.docx")
            data: Row data from Excel.
            index: Row index (0-based), used if pattern doesn't produce unique name.

        Returns:
            Generated filename.
        """
        filename = pattern

        for match in self.PLACEHOLDER_PATTERN.finditer(pattern):
            placeholder_name = match.group(1)
            value = data.get(placeholder_name, "")
            if value is None:
                value = ""
            # Sanitize filename
            value_str = str(value)
            for char in '<>:"/\\|?*':
                value_str = value_str.replace(char, "_")
            filename = filename.replace(match.group(0), value_str)

        # Ensure .docx extension
        if not filename.lower().endswith(".docx"):
            filename += ".docx"

        return filename
